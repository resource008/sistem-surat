from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DOCX_PATH = r"D:\sistem-surat\Laporan_Teknis_Magang_Tolan_Tiga_revisi_teknis.docx"


def insert_paragraph_after(anchor, text="", style=None):
    paragraph = anchor._parent.add_paragraph()
    anchor._p.addnext(paragraph._p)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def set_update_fields_on_open(document):
    settings = document.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def replace_heading_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


doc = Document(DOCX_PATH)

# Normalize a few important headings so the report reads as a technical report.
for p in doc.paragraphs:
    clean = " ".join(p.text.split())
    if clean == "Latar Belakang":
        replace_heading_text(p, "1.1 Latar Belakang")
    elif clean == "Analisis Sistem":
        replace_heading_text(p, "2.1 Analisis Sistem")
    elif clean == "Perancangan Sistem":
        replace_heading_text(p, "2.2 Perancangan Sistem")
    elif clean == "Lingkungan Pengembangan":
        replace_heading_text(p, "3.1 Lingkungan Pengembangan")
    elif clean == "Implementasi Backend dan Autentikasi":
        replace_heading_text(p, "3.2 Implementasi Backend dan Autentikasi")
    elif clean == "Implementasi Model Database (Prisma Schema)":
        replace_heading_text(p, "3.3 Implementasi Model Database (Prisma Schema)")
    elif clean == "Implementasi Route":
        replace_heading_text(p, "3.4 Implementasi Route")
    elif clean == "Dokumentasi API Endpoint":
        replace_heading_text(p, "3.5 Dokumentasi API Endpoint dan Pengujian Fungsional")
    elif clean == "Hasil Website":
        replace_heading_text(p, "3.6 Hasil Website")
    elif clean == "BAB IV KESIMPULAN":
        replace_heading_text(p, "BAB IV KESIMPULAN DAN SARAN")


# Strengthen Chapter I with the standard technical-report framing.
paras = doc.paragraphs
bab_ii = next(p for p in paras if " ".join(p.text.split()) == "BAB II ANALISIS DAN PERANCANGAN")
anchor = bab_ii._p.getprevious()
anchor_para = next(p for p in paras if p._p is anchor)

chapter_i_sections = [
    ("1.2 Rumusan Masalah", [
        "Berdasarkan kondisi pengelolaan surat masuk yang masih bergantung pada pencatatan manual dan lembar kerja Excel, rumusan masalah dalam laporan teknis ini adalah bagaimana merancang dan mengimplementasikan sistem registrasi surat masuk yang mampu mencatat data surat secara terpusat, mengatur hak akses pengguna, serta mendukung proses pelacakan status surat secara lebih cepat dan terdokumentasi.",
    ]),
    ("1.3 Tujuan", [
        "Tujuan pengembangan sistem ini adalah membangun aplikasi berbasis web untuk membantu proses registrasi, pencarian, pencetakan, dan pelacakan surat masuk di PT Tolan Tiga Indonesia.",
        "Secara khusus, sistem diarahkan untuk menyediakan autentikasi pengguna, pengelolaan data surat, pengelolaan akun, pengelolaan departemen, penomoran registrasi, serta fasilitas pelacakan yang dapat digunakan sesuai peran masing-masing pengguna.",
    ]),
    ("1.4 Batasan Masalah", [
        "Agar pembahasan laporan tetap terarah, ruang lingkup sistem dibatasi pada pengelolaan surat masuk, bukan seluruh siklus surat keluar atau arsip perusahaan secara menyeluruh.",
        "Sistem berfokus pada pengguna internal dengan pembagian peran Admin, Staf/PKL, dan Regional User. Pengujian yang dibahas dalam laporan ini difokuskan pada pengujian fungsional terhadap alur utama sistem dan endpoint API yang mendukung proses operasional.",
    ]),
    ("1.5 Manfaat", [
        "Manfaat yang diharapkan dari sistem ini adalah meningkatnya kerapian pencatatan surat masuk, berkurangnya risiko kehilangan data, tersedianya riwayat pelacakan yang lebih mudah dipantau, serta meningkatnya efisiensi kerja staf administrasi dalam melakukan pencarian dan pencetakan data surat.",
    ]),
    ("1.6 Metodologi Pelaksanaan", [
        "Penyusunan dan pengembangan sistem dilakukan melalui beberapa tahapan, yaitu observasi proses administrasi surat, diskusi kebutuhan dengan supervisor, analisis kebutuhan pengguna, perancangan alur kerja dan antarmuka, implementasi aplikasi web, serta pengujian fungsional berdasarkan skenario penggunaan utama.",
    ]),
]

for heading, paragraphs in chapter_i_sections:
    anchor_para = insert_paragraph_after(anchor_para, heading, "Heading 2")
    for text in paragraphs:
        anchor_para = insert_paragraph_after(anchor_para, text, "Normal")


# Add testing-method explanation before the API endpoint tables.
api_heading = next(
    p for p in doc.paragraphs
    if " ".join(p.text.split()) == "3.5 Dokumentasi API Endpoint dan Pengujian Fungsional"
)
testing_intro = [
    "Pengujian fungsional dilakukan dengan pendekatan black-box, yaitu memeriksa kesesuaian keluaran sistem terhadap masukan dan aksi pengguna tanpa membahas detail internal kode pada saat pengujian. Skenario pengujian disusun berdasarkan fungsi utama sistem, meliputi autentikasi, pengelolaan data surat, pencetakan surat, pengelolaan pengguna, dan pengelolaan departemen.",
    "Setiap endpoint diuji berdasarkan URL, metode HTTP, payload, dan respons yang diharapkan. Hasil pengujian digunakan untuk memastikan bahwa fitur utama berjalan sesuai kebutuhan operasional dan tidak hanya selesai dari sisi antarmuka.",
]
anchor_para = api_heading
for text in testing_intro:
    anchor_para = insert_paragraph_after(anchor_para, text, "Normal")


# Add a concise testing conclusion before section 3.6.
hasil_heading = next(p for p in doc.paragraphs if " ".join(p.text.split()) == "3.6 Hasil Website")
anchor_para = hasil_heading.insert_paragraph_before("Ringkasan Hasil Pengujian", "Heading 3")
anchor_para = insert_paragraph_after(
    anchor_para,
    "Berdasarkan skenario yang telah diuji, fungsi utama sistem dapat berjalan sesuai kebutuhan yang ditetapkan. Sistem mampu memvalidasi kredensial pengguna, menyimpan dan memperbarui data surat, mengambil detail surat, menghasilkan nomor registrasi berikutnya, menyiapkan data untuk proses cetak, serta mengelola data pengguna dan departemen melalui hak akses administrator.",
    "Normal",
)
anchor_para = insert_paragraph_after(
    anchor_para,
    "Temuan pengujian menunjukkan bahwa rancangan berbasis role membantu membatasi akses sesuai tanggung jawab pengguna. Dengan demikian, sistem tidak hanya mendukung pencatatan surat masuk, tetapi juga membantu menjaga konsistensi proses administrasi dan keterlacakan data.",
    "Normal",
)


# Fill Chapter IV with proper technical-report closing content.
bab_iv = next(
    p for p in doc.paragraphs
    if " ".join(p.text.split()) == "BAB IV KESIMPULAN DAN SARAN"
)
anchor_para = bab_iv
closing_sections = [
    ("4.1 Kesimpulan", [
        "Berdasarkan hasil analisis, perancangan, implementasi, dan pengujian yang telah dilakukan, Sistem Registrasi Surat Masuk dan Lacak Surat berhasil dikembangkan untuk mendukung proses administrasi surat di PT Tolan Tiga Indonesia. Sistem ini membantu mengubah proses pencatatan yang sebelumnya bergantung pada dokumen fisik dan lembar kerja manual menjadi proses digital yang lebih terstruktur.",
        "Fitur utama yang berhasil diimplementasikan meliputi autentikasi pengguna, pengelolaan data surat, pencetakan surat, pelacakan status surat, pengelolaan pengguna, dan pengelolaan departemen. Pembagian peran Admin, Staf/PKL, dan Regional User membuat penggunaan sistem lebih sesuai dengan kebutuhan operasional masing-masing pengguna.",
        "Pengujian fungsional terhadap endpoint dan alur utama menunjukkan bahwa sistem dapat menjalankan proses inti sesuai kebutuhan, mulai dari login, penambahan data, pengubahan data, penghapusan data, pengambilan data, hingga pengelolaan data master. Dengan demikian, sistem ini layak digunakan sebagai solusi pendukung administrasi surat masuk di lingkungan perusahaan.",
    ]),
    ("4.2 Saran", [
        "Pengembangan selanjutnya dapat diarahkan pada penambahan fitur notifikasi otomatis agar pengguna memperoleh pemberitahuan ketika status surat berubah atau ketika terdapat surat baru yang perlu ditindaklanjuti.",
        "Sistem juga dapat dikembangkan dengan fitur unggah dan arsip berkas digital agar dokumen fisik memiliki salinan elektronik yang tersimpan langsung bersama data registrasi surat.",
        "Selain itu, pengujian dapat diperluas dengan pengujian keamanan, pengujian performa, dan pengujian penerimaan pengguna agar sistem semakin siap digunakan dalam lingkungan kerja perusahaan secara berkelanjutan.",
    ]),
]

for heading, paragraphs in closing_sections:
    anchor_para = insert_paragraph_after(anchor_para, heading, "Heading 2")
    for text in paragraphs:
        anchor_para = insert_paragraph_after(anchor_para, text, "Normal")


# Use consistent alignment for inserted report prose.
for p in doc.paragraphs:
    style_name = p.style.name if p.style else ""
    if style_name == "Normal" and len(p.text.strip()) > 80:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

set_update_fields_on_open(doc)
doc.save(DOCX_PATH)
