from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("Tabel Aktivitas Februari-Juli 2026 Format Matriks.docx")

MONTHS = [
    {
        "month": "Februari 2026",
        "weeks": 3,
        "rows": [
            ("Mempelajari prosedur registrasi surat masuk secara manual", [1]),
            ("Pencatatan dan registrasi data surat masuk menggunakan Microsoft Excel", [1]),
            ("Analisis kekurangan dari Sistem Registrasi Surat Masuk dengan Excel", [1]),
            ("Rancangan desain antarmuka Sistem Registrasi Surat Masuk berbasis Website", [2]),
            ("Pembuatan alur sistem dengan Role-Based Access Control (RBAC)", [2]),
            ("Diskusi hasil rancangan desain sistem bersama Supervisor", [3]),
        ],
    },
    {
        "month": "Maret 2026",
        "weeks": 4,
        "rows": [
            ("Inisialisasi struktur awal proyek Sistem Registrasi Surat", [1]),
            ("Pembuatan autentikasi login menggunakan username dan password", [1]),
            ("Pengaturan session dan permission pengguna", [1]),
            ("Pengujian API login dan profil akun menggunakan Postman", [1]),
            ("Pembuatan halaman login berdasarkan role pengguna", [2]),
            ("Penyusunan alur Admin menuju halaman dashboard", [2]),
            ("Pengembangan halaman data surat untuk Staff/PKL", [3]),
            ("Pembuatan fitur filter tanggal dan departemen", [3]),
            ("Penyiapan route halaman cetak dan track surat", [3]),
            ("Pengembangan halaman detail data surat", [4]),
            ("Pengembangan form edit data surat", [4]),
            ("Pengujian API detail, edit, dan hapus data surat", [4]),
        ],
    },
    {
        "month": "April 2026",
        "weeks": 4,
        "rows": [
            ("Pengembangan form tambah surat untuk Staff/PKL", [1]),
            ("Integrasi pilihan departemen aktif pada form surat", [1]),
            ("Pengujian API preview nomor registrasi otomatis", [1]),
            ("Pengujian API tambah data surat baru", [1]),
            ("Penyempurnaan nomor register otomatis berdasarkan departemen dan tahun", [2]),
            ("Validasi input pada form tambah surat", [2]),
            ("Pengembangan halaman cetak surat sebagai bukti tanda terima", [3]),
            ("Pembuatan tampilan preview print cetak surat", [3]),
            ("Penyempurnaan pencarian data surat", [4]),
            ("Penyempurnaan filter tanggal surat, tanggal terima, perihal, dan departemen", [4]),
        ],
    },
    {
        "month": "Mei 2026",
        "weeks": 4,
        "rows": [
            ("Pengembangan halaman profil pengguna", [1]),
            ("Pengecekan session untuk akses halaman profil", [1]),
            ("Pengembangan halaman dashboard Admin", [2]),
            ("Integrasi data dashboard dari API", [2]),
            ("Pengembangan halaman Kelola Akun untuk Admin", [3]),
            ("Pengembangan halaman Kelola Role untuk Admin", [3]),
            ("Pengujian API Kelola Akun dan Kelola Role", [3]),
            ("Pengembangan fitur tambah akun pengguna", [4]),
            ("Pengembangan fitur edit akun pengguna", [4]),
            ("Validasi input akun dan pengaturan hak akses", [4]),
        ],
    },
    {
        "month": "Juni 2026",
        "weeks": 4,
        "rows": [
            ("Pengembangan fitur Kelola Departemen", [1]),
            ("Pembuatan data master departemen untuk data surat, filter, dan cetak surat", [1]),
            ("Pembuatan fitur tambah, edit, detail, sembunyikan, tampilkan, dan hapus departemen", [2]),
            ("Pengaturan kolom departemen berdasarkan tipe data dan kebutuhan tampilan", [2]),
            ("Pengujian API tambah, edit, hapus, dan detail departemen", [2]),
            ("Demonstrasi hasil pengembangan website kepada Supervisor", [3]),
            ("Peninjauan fitur dan pencatatan masukan dari Supervisor", [3]),
            ("Pengembangan backend fitur Kelola Sheet Lacak", [4]),
            ("Pengujian API daftar, tambah, detail, edit, urutan, sembunyikan, dan hapus Sheet Lacak", [4]),
        ],
    },
    {
        "month": "Juli 2026",
        "weeks": 2,
        "rows": [
            ("Penerapan data Sheet Lacak ke halaman Track Surat", [1]),
            ("Menampilkan kategori dan field berdasarkan Sheet Lacak", [1]),
            ("Pembuatan tampilan awal halaman Track Surat", [1]),
            ("Pembuatan proses tambah data pada halaman Track Surat", [2]),
            ("Pembuatan proses edit data pada halaman Track Surat", [2]),
            ("Penerapan hak akses field berdasarkan role pengguna", [2]),
        ],
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, italic_words=None, size=12, align=None):
    italic_words = italic_words or []
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.3
    if align is not None:
        paragraph.alignment = align

    chunks = [(text, False)]
    for word in italic_words:
        new_chunks = []
        for chunk, is_italic in chunks:
            if is_italic or word not in chunk:
                new_chunks.append((chunk, is_italic))
                continue
            before, match, after = chunk.partition(word)
            if before:
                new_chunks.append((before, False))
            new_chunks.append((match, True))
            if after:
                new_chunks.append((after, False))
        chunks = new_chunks

    for chunk, is_italic in chunks:
        run = paragraph.add_run(chunk)
        run.bold = bold
        run.italic = is_italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))


def make_caption(doc, number, month):
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    left = caption.add_run(f"Tabel {number}")
    left.bold = True
    left.font.name = "Times New Roman"
    left.font.size = Pt(12)
    spacer = caption.add_run("        ")
    spacer.font.name = "Times New Roman"
    right = caption.add_run(f"Tabel Aktivitas Bulan {month}")
    right.bold = True
    right.font.name = "Times New Roman"
    right.font.size = Pt(12)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.1)
section.right_margin = Inches(1.1)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(10)
title_run = title.add_run("Tabel Aktivitas Bulan Februari sampai Juli 2026")
title_run.bold = True
title_run.font.name = "Times New Roman"
title_run.font.size = Pt(14)

for table_number, month_data in enumerate(MONTHS, start=1):
    if table_number > 1:
        doc.add_paragraph()
    make_caption(doc, table_number, month_data["month"])

    weeks = month_data["weeks"]
    table = doc.add_table(rows=1, cols=1 + weeks)
    table.autofit = False
    set_table_borders(table)

    activity_width = Inches(2.65 if weeks == 4 else 2.45)
    week_width = Inches(0.95 if weeks == 4 else 1.18)
    total_width = int((activity_width.inches + (week_width.inches * weeks)) * 1440)
    set_table_width(table, total_width)

    header = table.rows[0].cells
    set_cell_width(header[0], activity_width)
    set_cell_text(header[0], "Aktivitas", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_margins(header[0])
    for index in range(weeks):
        cell = header[index + 1]
        set_cell_width(cell, week_width)
        set_cell_text(cell, f"Minggu ke-{index + 1}", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_margins(cell)

    for activity, active_weeks in month_data["rows"]:
        row = table.add_row().cells
        set_cell_width(row[0], activity_width)
        set_cell_text(
            row[0],
            activity,
            italic_words=["Microsoft Excel", "Excel", "Website", "Role-Based Access Control", "Supervisor", "Sheet Lacak", "Track Surat"],
            size=12,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )
        set_cell_margins(row[0], top=120, bottom=120)
        for week_index in range(1, weeks + 1):
            cell = row[week_index]
            set_cell_width(cell, week_width)
            set_cell_text(cell, "", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_margins(cell)
            if week_index in active_weeks:
                set_cell_shading(cell, "000000")

doc.save(OUTPUT)
print(OUTPUT.resolve())
